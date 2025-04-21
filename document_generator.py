"""
Module for generating document files from Wikipedia articles
"""

import os
import logging
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document_from_article(article, language):
    """
    Create a Word document from a Wikipedia article
    
    Args:
        article (dict): Article content dictionary
        language (str): Language code
        
    Returns:
        str: Path to the generated document file
    """
    if not article:
        return None
    
    try:
        # Create a new Document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = article['title']
        doc.core_properties.subject = f"Wikipedia article in {language}"
        
        # Add title
        title = doc.add_heading(article['title'], level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add language and source info
        source_paragraph = doc.add_paragraph()
        source_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        source_run = source_paragraph.add_run(f"Source: Wikipedia ({language})")
        source_run.italic = True
        
        # Add URL reference
        url_paragraph = doc.add_paragraph()
        url_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        url_run = url_paragraph.add_run(article['url'])
        url_run.italic = True
        
        # Add horizontal line
        doc.add_paragraph('_' * 50)
        
        # Add summary section
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(article['summary'])
        
        # Add full content
        doc.add_heading('Full Content', level=1)
        
        # Add sections if available
        if 'sections' in article and article['sections']:
            for section in article['sections']:
                if section.get('title'):
                    # Add section heading (level based on the section's level)
                    level = min(section.get('level', 0) + 2, 9)  # Limit to valid levels (Word only supports up to 9)
                    doc.add_heading(section['title'], level=level)
                
                # Add section content
                doc.add_paragraph(section['content'])
        else:
            # If no sections, just add the full content
            doc.add_paragraph(article['content'])
        
        # Create a safe filename
        safe_title = "".join([c if c.isalnum() or c in [' ', '-', '_'] else '_' for c in article['title']])
        filename = f"{safe_title}_{language}.docx"
        
        # Save the document to a temporary file
        file_path = os.path.join(os.getcwd(), filename)
        doc.save(file_path)
        
        return file_path
    
    except Exception as e:
        logging.error(f"Error creating document: {str(e)}")
        return None
